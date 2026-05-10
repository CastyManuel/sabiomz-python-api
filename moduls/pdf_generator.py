#!/usr/bin/env python3
"""
Sabiomz v3 — PDF & DOCX Generator
Gera PDF e DOCX profissional de um trabalho académico.
Uso: python3 pdf_generator.py <ficheiro_txt> <output_base> <titulo> <tarefa_id>
"""
import sys
import json
import os
import re


def clean_md(text: str) -> str:
    """Remove marcação Markdown simples para texto limpo."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)     # bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)           # italic
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.M) # headings
    text = re.sub(r'^---$', '', text, flags=re.M)       # hr
    return text.strip()


def gerar_pdf(texto: str, output_path: str, titulo: str) -> str | None:
    """Gera PDF com ReportLab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            HRFlowable, PageBreak, Table, TableStyle
        )
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        pdf_path = output_path + '.pdf'
        doc      = SimpleDocTemplate(
            pdf_path,
            pagesize=A4,
            leftMargin=3*cm, rightMargin=2.5*cm,
            topMargin=3*cm, bottomMargin=2.5*cm,
            title=titulo,
            author="Sabiomz AI",
        )

        styles = getSampleStyleSheet()

        # Estilos customizados
        st_capa = ParagraphStyle('Capa',
            fontName='Helvetica-Bold', fontSize=22, spaceAfter=20,
            alignment=TA_CENTER, textColor=colors.HexColor('#0A1F44'))
        st_sub = ParagraphStyle('SubCapa',
            fontName='Helvetica', fontSize=12, spaceAfter=8,
            alignment=TA_CENTER, textColor=colors.HexColor('#5b6b8a'))
        st_h1 = ParagraphStyle('H1',
            fontName='Helvetica-Bold', fontSize=14, spaceAfter=10, spaceBefore=20,
            textColor=colors.HexColor('#0A1F44'), borderPadding=(0,0,4,0))
        st_h2 = ParagraphStyle('H2',
            fontName='Helvetica-Bold', fontSize=12, spaceAfter=8, spaceBefore=12,
            textColor=colors.HexColor('#185FA5'))
        st_h3 = ParagraphStyle('H3',
            fontName='Helvetica-BoldOblique', fontSize=11, spaceAfter=6, spaceBefore=8,
            textColor=colors.HexColor('#00a86b'))
        st_body = ParagraphStyle('Body',
            fontName='Helvetica', fontSize=11, spaceAfter=6, spaceBefore=2,
            alignment=TA_JUSTIFY, leading=16)
        st_footer_txt = ParagraphStyle('Footer',
            fontName='Helvetica', fontSize=8, textColor=colors.grey, alignment=TA_CENTER)

        story = []

        # ── CAPA ──────────────────────────────────────────────────────────
        story.append(Spacer(1, 4*cm))
        story.append(Paragraph("REPÚBLICA DE MOÇAMBIQUE", st_sub))
        story.append(Spacer(1, 0.5*cm))
        story.append(HRFlowable(width="80%", color=colors.HexColor('#00a86b'), thickness=2, spaceBefore=6, spaceAfter=20))
        story.append(Paragraph(titulo.upper(), st_capa))
        story.append(HRFlowable(width="80%", color=colors.HexColor('#00a86b'), thickness=2, spaceBefore=6, spaceAfter=20))
        story.append(Spacer(1, 1*cm))
        story.append(Paragraph("Gerado por Sabiomz AI — www.sabiomz.co.mz", st_sub))
        story.append(Spacer(1, 6*cm))

        import datetime
        story.append(Paragraph(f"Maputo, {datetime.datetime.now().strftime('%B de %Y')}", st_sub))
        story.append(PageBreak())

        # ── CORPO ────────────────────────────────────────────────────────
        for line in texto.split('\n'):
            line = line.rstrip()
            if not line:
                story.append(Spacer(1, 4))
                continue

            # Headings
            if line.startswith('# '):
                story.append(Paragraph(clean_md(line[2:]), st_h1))
                story.append(HRFlowable(width="100%", color=colors.HexColor('#e0e8f5'), thickness=0.8, spaceAfter=6))
            elif line.startswith('## '):
                story.append(Paragraph(clean_md(line[3:]), st_h2))
            elif line.startswith('### '):
                story.append(Paragraph(clean_md(line[4:]), st_h3))
            elif line.startswith('---'):
                story.append(HRFlowable(width="100%", color=colors.lightgrey, thickness=0.5, spaceAfter=8))
            else:
                txt = clean_md(line)
                if txt:
                    # Bold inline
                    txt_html = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
                    txt_html = re.sub(r'\*(.+?)\*', r'<i>\1</i>', txt_html)
                    txt_html = re.sub(r'^#{1,6}\s+', '', txt_html)
                    try:
                        story.append(Paragraph(txt_html, st_body))
                    except Exception:
                        story.append(Paragraph(txt, st_body))

        # Adicionar rodapé no documento via onPage
        def footer(canvas, doc):
            canvas.saveState()
            canvas.setFont('Helvetica', 8)
            canvas.setFillColor(colors.grey)
            canvas.drawCentredString(A4[0]/2, 1.5*cm, f"Sabiomz AI  |  {titulo}  |  Pág. {doc.page}")
            canvas.restoreState()

        doc.build(story, onFirstPage=footer, onLaterPages=footer)
        return pdf_path

    except ImportError:
        return None
    except Exception as e:
        print(f"PDF error: {e}", file=sys.stderr)
        return None


def gerar_docx(texto: str, output_path: str, titulo: str) -> str | None:
    """Gera DOCX com python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        docx_path = output_path + '.docx'
        doc       = Document()

        # Margens
        for section in doc.sections:
            section.top_margin    = Cm(3)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(2.5)

        # ── CAPA ──────────────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("REPÚBLICA DE MOÇAMBIQUE\n")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x5b, 0x6b, 0x8a)

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titulo.upper())
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x0A, 0x1F, 0x44)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Gerado por Sabiomz AI\nwww.sabiomz.co.mz")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x5b, 0x6b, 0x8a)

        doc.add_page_break()

        # ── CORPO ────────────────────────────────────────────────────────
        for line in texto.split('\n'):
            line = line.rstrip()
            if not line:
                doc.add_paragraph()
                continue

            if line.startswith('# '):
                h = doc.add_heading(clean_md(line[2:]), level=1)
                h.style.font.color.rgb = RGBColor(0x0A, 0x1F, 0x44)
            elif line.startswith('## '):
                doc.add_heading(clean_md(line[3:]), level=2)
            elif line.startswith('### '):
                doc.add_heading(clean_md(line[4:]), level=3)
            elif line.startswith('---'):
                doc.add_paragraph('─' * 60)
            else:
                p   = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                txt = line
                # Bold inline **text**
                parts = re.split(r'\*\*(.+?)\*\*', txt)
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        p.add_run(part)
                    else:
                        run = p.add_run(part)
                        run.bold = True

        doc.save(docx_path)
        return docx_path

    except ImportError:
        return None
    except Exception as e:
        print(f"DOCX error: {e}", file=sys.stderr)
        return None


def main():
    if len(sys.argv) < 4:
        print(json.dumps({"erro": "Uso: pdf_generator.py <input.txt> <output_base> <titulo>"}))
        return

    input_file  = sys.argv[1]
    output_base = sys.argv[2]
    titulo      = sys.argv[3]

    if not os.path.exists(input_file):
        print(json.dumps({"erro": "Ficheiro de entrada não encontrado."}))
        return

    with open(input_file, 'r', encoding='utf-8') as f:
        texto = f.read()

    result = {}

    pdf_path = gerar_pdf(texto, output_base, titulo)
    if pdf_path:
        result['pdf'] = pdf_path
    else:
        result['pdf_erro'] = 'ReportLab não instalado. Execute: pip install reportlab'

    docx_path = gerar_docx(texto, output_base, titulo)
    if docx_path:
        result['docx'] = docx_path
    else:
        result['docx_erro'] = 'python-docx não instalado. Execute: pip install python-docx'

    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
