"""
Sabiomz v5 — modules/pdf_gen.py
Gera PDF profissional com ReportLab e DOCX com python-docx.
Fallback HTML quando as bibliotecas não estão disponíveis.
"""
import os
import re
import datetime


def _limpar_md(text: str) -> str:
    """Remove markdown simples para texto puro."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.M)
    text = re.sub(r"^---$", "", text, flags=re.M)
    return text.strip()


def gerar_pdf(texto: str, output_path: str, titulo: str) -> dict:
    """Gera PDF profissional. Retorna {'pdf': caminho} ou {'erro': msg, 'html': caminho_fallback}."""
    # Garantir directório
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            HRFlowable, PageBreak, KeepTogether
        )
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

        pdf_path = output_path if output_path.endswith(".pdf") else output_path + ".pdf"
        doc = SimpleDocTemplate(
            pdf_path, pagesize=A4,
            leftMargin=3*cm, rightMargin=2.5*cm,
            topMargin=3*cm, bottomMargin=2.5*cm,
            title=titulo, author="Sabiomz AI",
        )

        # Estilos
        cor_azul   = colors.HexColor("#0A1F44")
        cor_verde  = colors.HexColor("#00a86b")
        cor_muted  = colors.HexColor("#5b6b8a")

        st_capa  = ParagraphStyle("Capa",  fontName="Helvetica-Bold",   fontSize=22, spaceAfter=14, alignment=TA_CENTER, textColor=cor_azul)
        st_sub   = ParagraphStyle("Sub",   fontName="Helvetica",        fontSize=11, spaceAfter=6,  alignment=TA_CENTER, textColor=cor_muted)
        st_h1    = ParagraphStyle("H1",    fontName="Helvetica-Bold",   fontSize=14, spaceAfter=8,  spaceBefore=20, textColor=cor_azul)
        st_h2    = ParagraphStyle("H2",    fontName="Helvetica-Bold",   fontSize=12, spaceAfter=6,  spaceBefore=14, textColor=colors.HexColor("#185FA5"))
        st_h3    = ParagraphStyle("H3",    fontName="Helvetica-BoldOblique", fontSize=11, spaceAfter=5, spaceBefore=10, textColor=cor_verde)
        st_body  = ParagraphStyle("Body",  fontName="Helvetica",        fontSize=11, spaceAfter=5,  spaceBefore=2,  alignment=TA_JUSTIFY, leading=17)
        st_rodap = ParagraphStyle("Rodap", fontName="Helvetica",        fontSize=8,  textColor=colors.grey, alignment=TA_CENTER)

        story = []

        # ── CAPA ────────────────────────────────────────────
        story.append(Spacer(1, 4*cm))
        story.append(Paragraph("REPÚBLICA DE MOÇAMBIQUE", st_sub))
        story.append(HRFlowable(width="80%", color=cor_verde, thickness=2, spaceBefore=8, spaceAfter=18))
        story.append(Paragraph(titulo.upper(), st_capa))
        story.append(HRFlowable(width="80%", color=cor_verde, thickness=2, spaceBefore=8, spaceAfter=18))
        story.append(Spacer(1, 0.8*cm))
        story.append(Paragraph("Gerado por Sabiomz AI — www.sabiomz.co.mz", st_sub))
        story.append(Spacer(1, 4*cm))
        story.append(Paragraph(f"Maputo, {datetime.datetime.now().strftime('%B de %Y')}", st_sub))
        story.append(PageBreak())

        # ── CORPO ────────────────────────────────────────────
        for line in texto.split("\n"):
            line = line.rstrip()
            if not line:
                story.append(Spacer(1, 4))
                continue
            if line.startswith("# "):
                story.append(Paragraph(_limpar_md(line[2:]), st_h1))
                story.append(HRFlowable(width="100%", color=colors.HexColor("#e0e8f5"), thickness=0.7, spaceAfter=5))
            elif line.startswith("## "):
                story.append(Paragraph(_limpar_md(line[3:]), st_h2))
            elif line.startswith("### "):
                story.append(Paragraph(_limpar_md(line[4:]), st_h3))
            elif line.startswith("---"):
                story.append(HRFlowable(width="100%", color=colors.lightgrey, thickness=0.5, spaceAfter=6))
            else:
                txt_html = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)
                txt_html = re.sub(r"\*(.+?)\*", r"<i>\1</i>", txt_html)
                txt_html = re.sub(r"^#{1,6}\s+", "", txt_html)
                try:
                    story.append(Paragraph(txt_html, st_body))
                except Exception:
                    story.append(Paragraph(_limpar_md(line), st_body))

        def footer(canvas, doc):
            canvas.saveState()
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(colors.grey)
            canvas.drawCentredString(A4[0]/2, 1.5*cm, f"Sabiomz AI  |  {titulo}  |  Pág. {doc.page}")
            canvas.restoreState()

        doc.build(story, onFirstPage=footer, onLaterPages=footer)
        return {"pdf": pdf_path}

    except ImportError:
        # Fallback HTML
        html_path = (output_path if output_path.endswith(".html") else output_path + ".html")
        html_result = _gerar_html_fallback(texto, html_path, titulo)
        return {"erro_pdf": "ReportLab não instalado (pip install reportlab)", **html_result}
    except Exception as e:
        html_path = output_path.replace(".pdf", ".html")
        html_result = _gerar_html_fallback(texto, html_path, titulo)
        return {"erro_pdf": str(e), **html_result}


def gerar_docx(texto: str, output_path: str, titulo: str) -> dict:
    """Gera DOCX profissional. Retorna {'docx': caminho} ou {'erro': msg}."""
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        docx_path = output_path if output_path.endswith(".docx") else output_path + ".docx"
        doc = Document()

        # Margens
        for section in doc.sections:
            section.top_margin    = Cm(3)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(2.5)

        # ── CAPA ────────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("REPÚBLICA DE MOÇAMBIQUE\n")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x5b, 0x6b, 0x8a)

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titulo.upper())
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x0A, 0x1F, 0x44)

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Gerado por Sabiomz AI\nwww.sabiomz.co.mz\n{datetime.datetime.now().strftime('%B de %Y')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x5b, 0x6b, 0x8a)

        doc.add_page_break()

        # ── CORPO ────────────────────────────────────────────
        for line in texto.split("\n"):
            line = line.rstrip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith("# "):
                h = doc.add_heading(_limpar_md(line[2:]), level=1)
                try:
                    h.runs[0].font.color.rgb = RGBColor(0x0A, 0x1F, 0x44)
                except Exception:
                    pass
            elif line.startswith("## "):
                doc.add_heading(_limpar_md(line[3:]), level=2)
            elif line.startswith("### "):
                doc.add_heading(_limpar_md(line[4:]), level=3)
            elif line.startswith("---"):
                doc.add_paragraph("─" * 55)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Bold inline **...**
                partes = re.split(r"\*\*(.+?)\*\*", line)
                for i, parte in enumerate(partes):
                    run = p.add_run(parte)
                    if i % 2 == 1:
                        run.bold = True

        doc.save(docx_path)
        return {"docx": docx_path}

    except ImportError:
        return {"erro_docx": "python-docx não instalado (pip install python-docx)"}
    except Exception as e:
        return {"erro_docx": str(e)}


def _gerar_html_fallback(texto: str, html_path: str, titulo: str) -> dict:
    """Gera HTML como fallback quando PDF/DOCX não está disponível."""
    try:
        html = f"""<!DOCTYPE html>
<html lang="pt">
<head>
<meta charset="UTF-8">
<title>{_esc(titulo)}</title>
<style>
  body{{font-family:Georgia,serif;max-width:800px;margin:40px auto;line-height:1.9;color:#222;padding:0 20px}}
  h1{{color:#0A1F44;border-bottom:3px solid #00a86b;padding-bottom:8px;margin-top:30px}}
  h2{{color:#185FA5;margin-top:28px}}
  h3{{color:#00a86b}}
  p{{text-align:justify;margin:10px 0}}
  hr{{border:none;border-top:1px solid #ddd;margin:20px 0}}
  .capa{{text-align:center;padding:60px 0 80px;border-bottom:1px solid #ddd;margin-bottom:40px}}
  .capa h1{{font-size:28px;border:none}}
  footer{{margin-top:60px;border-top:1px solid #ddd;padding-top:12px;font-size:12px;color:#999;text-align:center}}
  @media print{{body{{margin:20mm}}}}
</style>
</head>
<body>
<div class="capa">
  <p style="color:#5b6b8a;font-size:13px">REPÚBLICA DE MOÇAMBIQUE</p>
  <h1>{_esc(titulo)}</h1>
  <p style="color:#5b6b8a;font-size:12px">Gerado por Sabiomz AI — www.sabiomz.co.mz<br>{datetime.datetime.now().strftime('%B de %Y')}</p>
</div>
"""
        for line in texto.split("\n"):
            line = line.rstrip()
            if not line:
                html += "<br>"
                continue
            if line.startswith("# "):
                html += f"<h1>{_esc(line[2:])}</h1>\n"
            elif line.startswith("## "):
                html += f"<h2>{_esc(line[3:])}</h2>\n"
            elif line.startswith("### "):
                html += f"<h3>{_esc(line[4:])}</h3>\n"
            elif line.startswith("---"):
                html += "<hr>\n"
            else:
                txt = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", _esc(line))
                html += f"<p>{txt}</p>\n"

        html += f"""<footer>Sabiomz AI | {_esc(titulo)} | {datetime.datetime.now().strftime('%d/%m/%Y')}</footer>
</body></html>"""

        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html)
        return {"html": html_path}
    except Exception as e:
        return {"erro_html": str(e)}


def _esc(t: str) -> str:
    """HTML escape simples."""
    return (t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
             .replace('"', "&quot;").replace("'", "&#x27;"))
