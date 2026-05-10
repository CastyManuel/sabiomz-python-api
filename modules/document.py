"""
Sabiomz v5 — modules/document.py
Lê e extrai texto de PDF, DOCX e TXT.
"""
import os
import re


def ler_documento(path: str) -> dict:
    """Extrai texto de um ficheiro. Suporta PDF, DOCX, TXT."""
    if not os.path.exists(path):
        return {"erro": f"Ficheiro não encontrado: {path}", "texto": ""}

    ext = os.path.splitext(path)[1].lower()
    tamanho_kb = round(os.path.getsize(path) / 1024, 1)

    if ext == ".pdf":
        resultado = _ler_pdf(path)
    elif ext in [".docx", ".doc"]:
        resultado = _ler_docx(path)
    elif ext == ".txt":
        resultado = _ler_txt(path)
    else:
        resultado = {"erro": f"Formato não suportado: {ext}", "texto": ""}

    if resultado.get("texto"):
        texto = _limpar_texto(resultado["texto"])
        resultado["texto"] = texto[:50000]  # limite 50k chars
        resultado["num_chars"] = len(texto)
        resultado["num_palavras"] = len(texto.split())
        resultado["truncado"] = len(texto) > 50000

    resultado["ficheiro"] = os.path.basename(path)
    resultado["tamanho_kb"] = tamanho_kb
    return resultado


def _ler_pdf(path: str) -> dict:
    # Tentar pdfplumber (mais preciso)
    try:
        import pdfplumber
        texto = ""
        paginas = 0
        with pdfplumber.open(path) as pdf:
            paginas = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texto += t + "\n\n"
        if texto.strip():
            return {"texto": texto.strip(), "paginas": paginas, "metodo": "pdfplumber"}
    except ImportError:
        pass
    except Exception:
        pass

    # Fallback PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        paginas = len(reader.pages)
        texto = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                texto += t + "\n\n"
        if texto.strip():
            return {"texto": texto.strip(), "paginas": paginas, "metodo": "pypdf2"}
    except ImportError:
        pass
    except Exception:
        pass

    return {"erro": "Não foi possível extrair texto do PDF. Instala: pip install pdfplumber", "texto": "", "paginas": 0}


def _ler_docx(path: str) -> dict:
    try:
        from docx import Document
        doc = Document(path)
        texto = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return {"texto": texto.strip(), "paginas": max(1, len(doc.paragraphs) // 20), "metodo": "python-docx"}
    except ImportError:
        return {"erro": "python-docx não instalado. Executa: pip install python-docx", "texto": ""}
    except Exception as e:
        return {"erro": str(e), "texto": ""}


def _ler_txt(path: str) -> dict:
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            with open(path, "r", encoding=enc) as f:
                texto = f.read()
            linhas = texto.count("\n")
            return {"texto": texto, "paginas": max(1, linhas // 40), "metodo": "txt", "encoding": enc}
        except UnicodeDecodeError:
            continue
        except Exception as e:
            return {"erro": str(e), "texto": ""}
    return {"erro": "Não foi possível detectar o encoding do ficheiro.", "texto": ""}


def _limpar_texto(texto: str) -> str:
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    texto = re.sub(r" {2,}", " ", texto)
    return texto.strip()
