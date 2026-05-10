#!/usr/bin/env python3
"""
Sabiomz v3 — Document Reader
Extrai texto de PDF, DOCX, TXT.
Uso: python3 document_reader.py <caminho_do_ficheiro>
"""
import sys
import json
import os


def read_pdf(path: str) -> dict:
    """Extrai texto de um PDF."""
    texto = ""
    paginas = 0

    # Tentar pdfplumber primeiro (mais preciso)
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            paginas = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texto += t + "\n\n"
        if texto.strip():
            return {"texto": texto.strip(), "paginas": paginas, "metodo": "pdfplumber"}
    except ImportError:
        pass
    except Exception as e:
        pass

    # Fallback: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        paginas = len(reader.pages)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                texto += t + "\n\n"
        if texto.strip():
            return {"texto": texto.strip(), "paginas": paginas, "metodo": "pypdf2"}
    except ImportError:
        pass
    except Exception as e:
        pass

    return {"erro": "Não foi possível extrair texto do PDF.", "texto": "", "paginas": 0}


def read_docx(path: str) -> dict:
    """Extrai texto de um DOCX."""
    try:
        from docx import Document
        doc    = Document(path)
        texto  = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return {"texto": texto.strip(), "paginas": len(doc.paragraphs) // 20, "metodo": "python-docx"}
    except ImportError:
        return {"erro": "python-docx não instalado. Execute: pip install python-docx", "texto": ""}
    except Exception as e:
        return {"erro": str(e), "texto": ""}


def read_txt(path: str) -> dict:
    """Lê ficheiro de texto."""
    try:
        for enc in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
            try:
                with open(path, 'r', encoding=enc) as f:
                    texto = f.read()
                linhas = texto.count('\n')
                return {"texto": texto, "paginas": max(1, linhas // 40), "metodo": "txt", "encoding": enc}
            except UnicodeDecodeError:
                continue
        return {"erro": "Não foi possível detectar o encoding do ficheiro.", "texto": ""}
    except Exception as e:
        return {"erro": str(e), "texto": ""}


def limpar_texto(texto: str) -> str:
    """Remove linhas em branco excessivas e espaços desnecessários."""
    import re
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    texto = re.sub(r' {2,}', ' ', texto)
    return texto.strip()


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"erro": "Uso: document_reader.py <caminho>"}))
        return

    path = sys.argv[1]

    if not os.path.exists(path):
        print(json.dumps({"erro": f"Ficheiro não encontrado: {path}"}))
        return

    ext    = os.path.splitext(path)[1].lower()
    tamanho = os.path.getsize(path)

    if ext == '.pdf':
        result = read_pdf(path)
    elif ext in ['.docx', '.doc']:
        result = read_docx(path)
    elif ext == '.txt':
        result = read_txt(path)
    else:
        result = {"erro": f"Extensão não suportada: {ext}"}

    # Limpar e truncar (máx 50 000 chars para não estourar memória)
    if result.get("texto"):
        result["texto"]      = limpar_texto(result["texto"])
        result["num_chars"]  = len(result["texto"])
        result["num_palavras"] = len(result["texto"].split())
        if len(result["texto"]) > 50000:
            result["texto"]   = result["texto"][:50000]
            result["truncado"] = True

    result["ficheiro"] = os.path.basename(path)
    result["tamanho_kb"] = round(tamanho / 1024, 1)

    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()v
